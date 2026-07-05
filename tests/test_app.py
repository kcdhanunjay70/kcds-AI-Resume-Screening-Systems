from app import app


SAMPLE_RESUME = """
Python developer with 2 years experience building Flask REST API projects.
Worked on MongoDB database integration, SQL queries, Git workflow, testing,
deployment and backend debugging. Created college projects with OOPS concepts
and responsive HTML CSS JavaScript dashboards.
"""


def test_health_endpoint():
    client = app.test_client()
    response = client.get("/api/health")
    payload = response.get_json()
    assert response.status_code == 200
    assert payload["success"] is True


def test_metadata_endpoint_lists_roles():
    client = app.test_client()
    response = client.get("/api/metadata")
    payload = response.get_json()
    assert response.status_code == 200
    assert "Python Developer" in payload["roles"]
    assert "Full Stack Developer" in payload["roles"]
    assert len(payload["jobs"]) >= 30
    assert all(job["salaryMin"] < job["salaryMax"] for job in payload["jobs"])


def test_screen_resume_returns_score_and_history():
    client = app.test_client()
    response = client.post(
        "/api/screen",
        json={
            "candidateName": "Test Candidate",
            "email": "candidate@example.com",
            "phone": "9999999999",
            "role": "Python Developer",
            "minExperience": 1,
            "resumeText": SAMPLE_RESUME,
        },
    )
    payload = response.get_json()
    assert response.status_code == 201
    assert payload["screening"]["finalScore"] > 50
    assert "python" in payload["screening"]["matchedSkills"]
    assert payload["screening"]["salaryMax"] == 1400000
    assert payload["screening"]["classification"]["primaryDomain"]

    history = client.get("/api/screenings").get_json()
    assert any(item["candidateName"] == "Test Candidate" for item in history["screenings"])


def test_short_resume_validation():
    client = app.test_client()
    response = client.post("/api/screen", json={"role": "Python Developer", "resumeText": "too short"})
    assert response.status_code == 400
    assert response.get_json()["success"] is False


def test_txt_resume_extraction_and_classification():
    client = app.test_client()
    response = client.post(
        "/api/resume/extract",
        data={"resume": (io.BytesIO(SAMPLE_RESUME.encode()), "resume.txt")},
        content_type="multipart/form-data",
    )
    payload = response.get_json()
    assert response.status_code == 200
    assert "Python developer" in payload["text"]
    assert payload["classification"]["wordCount"] > 20
    assert payload["classification"]["recommendedJobs"]


def test_docx_resume_extraction():
    document = Document()
    document.add_heading("Software Engineer Resume", 0)
    document.add_paragraph(SAMPLE_RESUME)
    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    response = app.test_client().post(
        "/api/resume/extract",
        data={"resume": (stream, "resume.docx")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    assert "MongoDB" in response.get_json()["text"]


def test_rejects_unsupported_resume_file():
    response = app.test_client().post(
        "/api/resume/extract",
        data={"resume": (io.BytesIO(b"not a resume"), "resume.exe")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
import io

from docx import Document
