import io
import os
import re
from collections import Counter
from datetime import datetime, timezone
from uuid import uuid4

from bson import ObjectId
from docx import Document
from flask import Flask, jsonify, render_template, request
from pypdf import PdfReader
from pymongo import MongoClient
from pymongo.errors import PyMongoError


def job(job_id, title, category, skills, keywords, salary_min, salary_max, experience, level="Mid Level"):
    return {
        "id": job_id, "title": title, "category": category, "skills": skills,
        "keywords": keywords, "salaryMin": salary_min, "salaryMax": salary_max,
        "experience": experience, "level": level, "currency": "INR", "period": "year",
    }


JOB_CATALOG = [
    job("python-developer", "Python Developer", "Software Development", ["python", "flask", "django", "rest api", "mongodb", "sql", "git", "testing", "oops"], ["backend", "api", "database", "debugging", "deployment"], 500000, 1400000, 2),
    job("java-developer", "Java Developer", "Software Development", ["java", "spring boot", "hibernate", "rest api", "sql", "microservices", "git", "junit"], ["backend", "api", "database", "deployment"], 550000, 1600000, 2),
    job("dotnet-developer", ".NET Developer", "Software Development", ["c#", ".net", "asp.net", "sql server", "rest api", "azure", "git"], ["backend", "api", "enterprise", "database"], 550000, 1500000, 2),
    job("node-developer", "Node.js Developer", "Software Development", ["javascript", "node.js", "express", "mongodb", "rest api", "typescript", "git"], ["backend", "api", "async", "database"], 550000, 1500000, 2),
    job("frontend-developer", "Frontend Developer", "Web Engineering", ["html", "css", "javascript", "react", "responsive design", "api integration", "git"], ["ui", "components", "browser", "accessibility", "performance"], 450000, 1400000, 1),
    job("react-developer", "React Developer", "Web Engineering", ["react", "javascript", "typescript", "redux", "html", "css", "jest", "git"], ["components", "hooks", "state management", "frontend"], 550000, 1600000, 2),
    job("angular-developer", "Angular Developer", "Web Engineering", ["angular", "typescript", "rxjs", "html", "css", "rest api", "git"], ["components", "spa", "frontend", "testing"], 550000, 1550000, 2),
    job("full-stack-developer", "Full Stack Developer", "Web Engineering", ["html", "css", "javascript", "python", "flask", "mongodb", "rest api", "deployment", "git"], ["frontend", "backend", "authentication", "database", "testing"], 650000, 1900000, 2),
    job("mobile-developer", "Mobile App Developer", "Mobile", ["flutter", "dart", "android", "ios", "rest api", "firebase", "git"], ["mobile", "app", "deployment", "ui"], 500000, 1500000, 2),
    job("android-developer", "Android Developer", "Mobile", ["kotlin", "java", "android studio", "rest api", "firebase", "git"], ["mobile", "material design", "play store"], 500000, 1500000, 2),
    job("ios-developer", "iOS Developer", "Mobile", ["swift", "swiftui", "xcode", "rest api", "core data", "git"], ["mobile", "app store", "ios"], 600000, 1800000, 2),
    job("data-analyst", "Data Analyst", "Data & AI", ["python", "sql", "excel", "power bi", "tableau", "statistics", "pandas", "dashboard"], ["reporting", "insights", "visualization", "cleaning", "analysis"], 450000, 1200000, 1),
    job("business-analyst", "Business Analyst", "Data & AI", ["excel", "sql", "power bi", "requirements", "jira", "documentation", "analytics"], ["stakeholder", "process", "insights", "reporting"], 500000, 1400000, 2),
    job("data-scientist", "Data Scientist", "Data & AI", ["python", "machine learning", "statistics", "pandas", "numpy", "sql", "scikit-learn"], ["modeling", "experiments", "features", "prediction"], 700000, 2200000, 2),
    job("ml-engineer", "Machine Learning Engineer", "Data & AI", ["python", "machine learning", "scikit-learn", "pandas", "numpy", "model deployment", "mongodb"], ["classification", "regression", "training", "features", "accuracy"], 800000, 2400000, 2),
    job("ai-engineer", "AI Engineer", "Data & AI", ["python", "deep learning", "pytorch", "tensorflow", "llm", "nlp", "model deployment"], ["artificial intelligence", "training", "inference", "evaluation"], 900000, 2800000, 2),
    job("data-engineer", "Data Engineer", "Data & AI", ["python", "sql", "spark", "airflow", "etl", "aws", "data warehouse"], ["pipeline", "big data", "analytics", "cloud"], 700000, 2200000, 2),
    job("devops-engineer", "DevOps Engineer", "Cloud & Operations", ["linux", "docker", "kubernetes", "jenkins", "aws", "terraform", "git"], ["ci/cd", "deployment", "monitoring", "automation"], 700000, 2100000, 2),
    job("cloud-engineer", "Cloud Engineer", "Cloud & Operations", ["aws", "azure", "gcp", "linux", "networking", "terraform", "docker"], ["cloud", "infrastructure", "security", "deployment"], 700000, 2200000, 2),
    job("site-reliability-engineer", "Site Reliability Engineer", "Cloud & Operations", ["linux", "kubernetes", "prometheus", "grafana", "python", "cloud", "incident management"], ["reliability", "observability", "slo", "automation"], 900000, 2600000, 3),
    job("system-administrator", "System Administrator", "Cloud & Operations", ["linux", "windows server", "networking", "active directory", "powershell", "backup"], ["support", "infrastructure", "monitoring"], 400000, 1100000, 2),
    job("cybersecurity-analyst", "Cybersecurity Analyst", "Cybersecurity", ["network security", "siem", "linux", "incident response", "vulnerability assessment", "firewall"], ["security", "threat", "risk", "monitoring"], 600000, 1800000, 2),
    job("security-engineer", "Security Engineer", "Cybersecurity", ["application security", "cloud security", "penetration testing", "owasp", "siem", "python"], ["security", "vulnerability", "compliance", "threat"], 800000, 2400000, 3),
    job("qa-engineer", "QA Engineer", "Quality Engineering", ["manual testing", "test cases", "jira", "api testing", "sql", "agile"], ["quality", "defect", "regression", "validation"], 400000, 1100000, 1),
    job("automation-test-engineer", "Automation Test Engineer", "Quality Engineering", ["selenium", "python", "java", "api testing", "pytest", "jenkins", "git"], ["automation", "regression", "framework", "quality"], 550000, 1600000, 2),
    job("ui-ux-designer", "UI/UX Designer", "Product & Design", ["figma", "wireframes", "prototyping", "user research", "design system", "responsive design"], ["design", "usability", "user experience", "interaction"], 450000, 1400000, 1),
    job("product-manager", "Product Manager", "Product & Design", ["product strategy", "roadmap", "analytics", "jira", "agile", "user research"], ["product", "stakeholder", "metrics", "delivery"], 900000, 2800000, 3),
    job("scrum-master", "Scrum Master", "Product & Design", ["scrum", "agile", "jira", "sprint planning", "retrospective", "risk management"], ["facilitation", "delivery", "team", "process"], 700000, 1900000, 3),
    job("database-administrator", "Database Administrator", "Database", ["sql", "mysql", "postgresql", "mongodb", "backup", "performance tuning"], ["database", "availability", "security", "recovery"], 600000, 1800000, 2),
    job("mongodb-developer", "MongoDB Developer", "Database", ["mongodb", "aggregation", "indexing", "python", "node.js", "data modeling", "atlas"], ["nosql", "database", "performance", "api"], 650000, 1900000, 2),
    job("blockchain-developer", "Blockchain Developer", "Emerging Technology", ["solidity", "ethereum", "web3", "javascript", "smart contracts", "git"], ["blockchain", "defi", "security", "distributed"], 800000, 2500000, 2),
    job("technical-support-engineer", "Technical Support Engineer", "IT Support", ["troubleshooting", "linux", "windows", "networking", "ticketing", "customer support"], ["support", "incident", "communication", "resolution"], 350000, 900000, 1, "Entry Level"),
]

JOB_MAP = {item["id"]: item for item in JOB_CATALOG}
TITLE_MAP = {item["title"]: item for item in JOB_CATALOG}
EXPERIENCE_WORDS = {"internship": 5, "project": 4, "deployed": 5, "production": 5, "certification": 3, "lead": 4, "team": 2, "client": 3, "github": 4}
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MAX_FILE_BYTES = 8 * 1024 * 1024


def utc_now():
    return datetime.now(timezone.utc)


def clean(value, max_length=500):
    return str(value or "").strip()[:max_length]


def serialize(document):
    item = dict(document)
    if "_id" in item:
        item["id"] = str(item.pop("_id"))
    for key in ("createdAt", "updatedAt"):
        if isinstance(item.get(key), datetime):
            item[key] = item[key].isoformat()
    return item


def tokens(text):
    return re.findall(r"[a-z0-9+#.]+", str(text or "").lower())


def contains_phrase(text, phrase):
    return phrase.lower() in text.lower()


def resolve_job(value):
    value = clean(value, 100)
    selected = JOB_MAP.get(value) or TITLE_MAP.get(value)
    if not selected:
        raise ValueError("Please select a supported job")
    return selected


def extract_resume(file_storage):
    if not file_storage or not file_storage.filename:
        raise ValueError("Choose a TXT, PDF or DOCX resume")
    extension = os.path.splitext(file_storage.filename)[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Only TXT, PDF and DOCX resumes are supported")
    content = file_storage.read(MAX_FILE_BYTES + 1)
    if len(content) > MAX_FILE_BYTES:
        raise ValueError("Resume file must be 8 MB or smaller")
    try:
        if extension == ".pdf":
            text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)
        elif extension == ".docx":
            document = Document(io.BytesIO(content))
            text = "\n".join([p.text for p in document.paragraphs] + [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows])
        else:
            text = content.decode("utf-8-sig", errors="replace")
    except Exception as exc:
        raise ValueError(f"Could not read this {extension[1:].upper()} resume") from exc
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 80:
        raise ValueError("The document does not contain enough readable resume text")
    return text


def classify_resume(text):
    lower = text.lower()
    sections = [name for name in ["summary", "objective", "experience", "education", "skills", "projects", "certifications", "achievements"] if name in lower]
    years = [float(value) for value in re.findall(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year)", lower)]
    detected_experience = max(years) if years else 0
    category_scores = Counter()
    job_scores = []
    for item in JOB_CATALOG:
        skill_hits = [skill for skill in item["skills"] if contains_phrase(lower, skill)]
        keyword_hits = [word for word in item["keywords"] if contains_phrase(lower, word)]
        score = len(skill_hits) * 3 + len(keyword_hits)
        category_scores[item["category"]] += score
        job_scores.append((score, item))
    ranked = [item for score, item in sorted(job_scores, key=lambda pair: pair[0], reverse=True) if score > 0][:5]
    level = "Fresher / Entry Level" if detected_experience < 2 else "Mid Level" if detected_experience < 5 else "Senior Level"
    resume_type = "Professional Resume" if "experience" in sections and len(sections) >= 4 else "Project-based Resume" if "project" in sections else "General Resume"
    return {
        "resumeType": resume_type,
        "careerLevel": level,
        "detectedExperience": detected_experience,
        "primaryDomain": category_scores.most_common(1)[0][0] if category_scores else "General Software",
        "sectionsFound": sections,
        "wordCount": len(tokens(text)),
        "recommendedJobs": [{"id": item["id"], "title": item["title"], "salaryMin": item["salaryMin"], "salaryMax": item["salaryMax"]} for item in ranked],
    }


def score_resume(payload):
    candidate_name = clean(payload.get("candidateName"), 80) or "Candidate"
    selected_job = resolve_job(payload.get("jobId") or payload.get("role"))
    resume_text = clean(payload.get("resumeText"), 30000)
    custom_skills = clean(payload.get("requiredSkills"), 1000)
    try:
        min_experience = int(float(payload.get("minExperience") if payload.get("minExperience") not in (None, "") else selected_job["experience"]))
    except (TypeError, ValueError) as exc:
        raise ValueError("Minimum experience must be a valid number") from exc
    if len(resume_text) < 80:
        raise ValueError("Resume text must contain at least 80 characters")
    if not 0 <= min_experience <= 20:
        raise ValueError("Minimum experience must be between 0 and 20 years")

    required_skills = list(selected_job["skills"])
    for value in re.split(r"[,|\n]", custom_skills):
        skill = value.strip().lower()
        if skill and skill not in required_skills:
            required_skills.append(skill)
    resume_lower = resume_text.lower()
    word_counts = Counter(tokens(resume_text))
    matched = [skill for skill in required_skills if contains_phrase(resume_lower, skill)]
    missing = [skill for skill in required_skills if skill not in matched]
    keyword_hits = [word for word in selected_job["keywords"] if contains_phrase(resume_lower, word)]
    classification = classify_resume(resume_text)
    experience = classification["detectedExperience"]
    experience_score = 100 if experience >= min_experience else max(25, int((experience / max(min_experience, 1)) * 100))
    skill_score = int(len(matched) / max(len(required_skills), 1) * 100)
    keyword_score = int(len(keyword_hits) / max(len(selected_job["keywords"]), 1) * 100)
    quality_score = min(100, 38 + len(classification["sectionsFound"]) * 7 + len(resume_text) // 120 + sum(EXPERIENCE_WORDS.get(word, 0) for word in word_counts))
    final_score = round(skill_score * .54 + keyword_score * .16 + experience_score * .18 + quality_score * .12)
    if final_score >= 82:
        decision, recommendation = "Strong Match", "Shortlist for technical interview."
    elif final_score >= 65:
        decision, recommendation = "Good Match", "Review projects and conduct a screening call."
    elif final_score >= 48:
        decision, recommendation = "Potential Match", "Keep as backup and validate missing skills."
    else:
        decision, recommendation = "Needs Upskilling", "Build the missing skills before applying for this role."
    strengths = []
    if matched:
        strengths.append(f"Matched {len(matched)} of {len(required_skills)} required skills")
    if keyword_hits:
        strengths.append(f"Role evidence found: {', '.join(keyword_hits[:4])}")
    if experience:
        strengths.append(f"Detected {experience:g} years of experience")
    if classification["sectionsFound"]:
        strengths.append(f"Structured sections: {', '.join(classification['sectionsFound'][:5])}")
    return {
        "candidateName": candidate_name, "jobId": selected_job["id"], "role": selected_job["title"],
        "job": selected_job, "salaryMin": selected_job["salaryMin"], "salaryMax": selected_job["salaryMax"],
        "minExperience": min_experience, "detectedExperience": experience, "requiredSkills": required_skills,
        "matchedSkills": matched, "missingSkills": missing, "keywordHits": keyword_hits,
        "skillScore": skill_score, "keywordScore": keyword_score, "experienceScore": experience_score,
        "resumeQualityScore": quality_score, "finalScore": final_score, "decision": decision,
        "recommendation": recommendation, "strengths": strengths or ["Readable resume submitted"],
        "classification": classification, "resumePreview": resume_text[:900],
    }


class ScreeningStore:
    def __init__(self):
        self.client = self.db = None
        self.mode, self.error, self.memory_screenings = "memory", "", []
        mongo_uri = os.getenv("MONGO_URI") or os.getenv("MONGODB_URI")
        if mongo_uri:
            try:
                self.client = MongoClient(mongo_uri, serverSelectionTimeoutMS=2500)
                self.client.admin.command("ping")
                self.db = self.client[os.getenv("MONGO_DB_NAME", "ai_resume_screening")]
                self.mode = "mongodb"
                self.screenings.create_index("createdAt")
                self.screenings.create_index([("jobId", 1), ("finalScore", -1)])
            except PyMongoError as exc:
                self.error = str(exc)

    @property
    def screenings(self):
        return self.db["screenings"] if self.db is not None else None

    def create_screening(self, payload):
        result, now = score_resume(payload), utc_now()
        document = {**result, "email": clean(payload.get("email"), 120), "phone": clean(payload.get("phone"), 20), "notes": clean(payload.get("notes"), 500), "sourceFile": clean(payload.get("sourceFile"), 180), "createdAt": now, "updatedAt": now}
        if self.mode == "mongodb":
            document["_id"] = self.screenings.insert_one(document).inserted_id
        else:
            document["id"] = str(uuid4())
            self.memory_screenings.insert(0, document)
        return serialize(document)

    def list_screenings(self, limit=40):
        if self.mode == "mongodb":
            return [serialize(item) for item in self.screenings.find({}).sort("createdAt", -1).limit(limit)]
        return [serialize(item) for item in self.memory_screenings[:limit]]

    def delete_screening(self, screening_id):
        if self.mode == "mongodb":
            if not ObjectId.is_valid(screening_id) or not self.screenings.delete_one({"_id": ObjectId(screening_id)}).deleted_count:
                raise LookupError("Screening not found")
        else:
            before = len(self.memory_screenings)
            self.memory_screenings = [item for item in self.memory_screenings if item["id"] != screening_id]
            if before == len(self.memory_screenings):
                raise LookupError("Screening not found")

    def stats(self):
        items = self.list_screenings(500)
        return {
            "totalScreenings": len(items),
            "averageScore": round(sum(item["finalScore"] for item in items) / len(items), 1) if items else 0,
            "shortlisted": sum(item["decision"] in {"Strong Match", "Good Match"} for item in items),
            "roleCount": len({item["role"] for item in items}),
            "jobCount": len(JOB_CATALOG), "databaseReady": self.mode == "mongodb", "mode": self.mode,
            "databaseError": self.error,
        }


def create_app():
    app = Flask(__name__)
    app.config.update(JSON_SORT_KEYS=False, MAX_CONTENT_LENGTH=MAX_FILE_BYTES + 1024 * 1024)
    store = ScreeningStore()

    @app.get("/")
    def home():
        return render_template("index.html", jobs=JOB_CATALOG)

    @app.get("/api/health")
    def health():
        return jsonify({"success": True, "status": "ok", "storeMode": store.mode})

    @app.get("/api/metadata")
    def metadata():
        return jsonify({"success": True, "jobs": JOB_CATALOG, "roles": {item["title"]: item for item in JOB_CATALOG}})

    @app.get("/api/stats")
    def stats():
        return jsonify({"success": True, "stats": store.stats()})

    @app.get("/api/screenings")
    def screenings():
        return jsonify({"success": True, "screenings": store.list_screenings()})

    @app.post("/api/resume/extract")
    def resume_extract():
        try:
            uploaded = request.files.get("resume")
            text = extract_resume(uploaded)
            return jsonify({"success": True, "fileName": uploaded.filename, "text": text, "classification": classify_resume(text)})
        except ValueError as exc:
            return jsonify({"success": False, "message": str(exc)}), 400

    @app.post("/api/screen")
    def screen():
        try:
            payload = request.get_json(silent=True) or {}
            return jsonify({"success": True, "screening": store.create_screening(payload)}), 201
        except ValueError as exc:
            return jsonify({"success": False, "message": str(exc)}), 400
        except PyMongoError as exc:
            return jsonify({"success": False, "message": f"Database error: {exc}"}), 502

    @app.delete("/api/screenings/<screening_id>")
    def delete(screening_id):
        try:
            store.delete_screening(screening_id)
            return jsonify({"success": True})
        except LookupError as exc:
            return jsonify({"success": False, "message": str(exc)}), 404

    return app


app = create_app()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "5000")), debug=os.getenv("FLASK_DEBUG") == "1")
